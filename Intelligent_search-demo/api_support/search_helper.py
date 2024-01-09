from tika import parser
import uuid
import time
import os
import requests
from fastapi import FastAPI, Depends, HTTPException, Request
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from pydantic import BaseModel
from passlib.context import CryptContext
from datetime import datetime, timedelta
from jose import JWTError, jwt
from motor.motor_asyncio import AsyncIOMotorClient
from utils.config import *
from model.api_schema import * 
from fastapi import APIRouter,HTTPException
from utils.connection import search_obj
from datetime import datetime
import glob
from docx import Document
import xlrd

def process_docx(docx_file_location):
    doc = Document(docx_file_location)
    text_content = '/n'.join([paragraph.text for paragraph in doc.paragraphs])
    table_data = '/n'.join([''.join([cell.text for cell in row.cells]) for table in doc.tables for row in table.rows])
    data=text_content +'/n'+ table_data
    return {'content':data}


def process_pdf(pdf_file_location):
    " Extracts data from PDF, PPT file"
    data = parser.from_file(pdf_file_location)
    return data

#TO DO
def process_excel(excel_file_path):
    book = xlrd.open_workbook(excel_file_path)
    print("The number of worksheets is {0}".format(book.nsheets))
    print("Worksheet name(s): {0}".format(book.sheet_names()))
    sh = book.sheet_by_index(0)
    print("{0} {1} {2}".format(sh.name, sh.nrows, sh.ncols))
    print("Cell D30 is {0}".format(sh.cell_value(rowx=29, colx=3)))
    for rx in range(sh.nrows):
        print(sh.row(rx))

def load_dataset(file_path):
    file_list = os.listdir(file_path)

    count = 0
    for file in file_list:
        print("Processing file : {}, out of {} files.".format(count, len(file_list)))
        file_location = os.path.join(file_path, file)
        print("FILEEE LOCATION TO PRECESS", file_location)
        parsed_content={}
        #try:
        if file.lower().endswith('.docx') or file.lower().endswith('.doc'):
            parsed_content = process_docx(file_location)
            metadata={}
        else:
            parsed_content = process_pdf(file_location)
            try:
                metadata = {"created_on": parsed_content["metadata"]["dcterms:created"],
                            "modified_on" : parsed_content["metadata"]["dcterms:modified"],
                            "no_of_page" : parsed_content["metadata"]["xmpTPg:NPages"]}
            except:
                metadata={}
        # except Exception:
        #     print("Error in processing file")
        #     continue
        doc = {
            "filename": file,
            "parsed_content": parsed_content["content"],
            "pdf_metadata": metadata,
            "file_path":file_location,
            "added_by":"a8098c1a-f86e-11da-bd1a-00112444be1e",
            "added_on":datetime.now(),
            "doc_type":file.split('.')[-1]
        }

        resp = search_obj.index(
            index = INDEX_NAME,
            body = doc,
            id = uuid.uuid4(),
            refresh = True
        )
        print('\nAdding document:')
        print(resp)
    
    # refresh the index to make the documents searchable
    search_obj.indices.refresh(index=INDEX_NAME)
    return file_list